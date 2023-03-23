import pickle
import sklearn
from sanitybot.models import User_health
from flask_login import login_required, current_user
from sanitybot import db
from sanitybot.models import User
from datetime import datetime
from docx import Document
import docx
from io import BytesIO
from datetime import datetime

with open(r'sanitybot\pickles\PHQ9_decisiontree.pkl', 'rb') as f:
    decisionTree_model = pickle.load(f)
with open(r'sanitybot\pickles\PHQ9_naiveBayes.pkl', 'rb') as f:
    naiveBayes_model = pickle.load(f)
with open(r'sanitybot\pickles\model\PHQ9_lstm.pkl', 'rb') as f:
    lstm_model = pickle.load(f)


def get_frequency_label(frequency):
    switcher = {
        0: "Not at all",
        1: "Several days",
        2: "More than half the days",
        3: "Nearly everyday"
    }
    return switcher.get(frequency, "Invalid frequency")


def get_answers(phq9, input_lists):
    answers = []
    for i in range(len(phq9)):
        question = phq9[i]['question']
        answer = phq9[i][input_lists[i]]
        # answers.append((question, answer)) #question and answers
        answers.append(answer)
    return answers


def pdd_prediction(input_list):
    pred_result = lstm_model.predict([input_list])[0]
    print(pred_result)

    phq9 = [
        {
            "question": "Are you feeling despair when you are doing something that you typically enjoy??",
            0: "Nope, not at all.",
            1: "Yes, I feel that way",
            2: "Almost an entire day",
            3: "Consistently day after day",
        },
        {
            "question": "I've noticed that you've been isolating yourself and feeling down more lately and may be experiencing feelings of hopelessness. Have you noticed any changes or improvements in your mood and outlook?",
            0: "No, I dont noticed any changes",
            1: "Yes, I experiecing that kind of feeling.",
            2: "Sometimes I feel that way and I noticed it.",
            3: "I'm worried i felt that way every day.",
        },
        {
            "question": "But how was your sleeping schedule lately? How was your sleep been, and have you been experiencing any difficulty falling or staying asleep? You seem like you're going through a tough time.",
            0: "My sleep schedule is okay.",
            1: "Its not good lately",
            2: "Its hard for me to get a sleep.",
            3: "I'm having difficulty getting to sleep.",
        },
        {
            "question": "It will appears in feeling tired and low on energy. So, have you noticed any patterns or triggers that might be contributing to your fatigue like particular reasons for your lack of energy, such as stress or changes in your routine?",
            0: "Not even a little bit",
            1: "Yes, sometimes.",
            2: "Half of the day I almost faint",
            3: "What to do if I feel fatigue everday",
        },
        {
            "question": "Oh I see. Somehow I sense that you may be experiencing feelings of shame or self-blame. Is there anything thats been weighing heavily on your mind, such as past mistakes or current challenges that can me you feel of failure or disappointment anything in particular that has been contributing to these thoughts?",
            0: "No, I dont have any thoughts like that.",
            1: "Yes, sometimes I feel anxious about that.",
            2: "I often find myself dwelling on the mistakes I've made before.",
            3: "My past mistakes still weigh heavily on my mind.",
        },
        {
            "question": "Oh, but one thing that might help to avoid bad thoughts is to focus. So, how is your experience uncontrollable worrying thoughts that seem to dominate your mind, and does it affect your daily life and activities?",
            0: "Im not experiencing any of those",
            1: "Quite a few days I guess",
            2: "I spend about half the day feeling that way.",
            3: "I'm anxious about experiencing this feeling every day",
        },
        {
            "question": "I understand. But have you noticed that you frequently have an overwhelming sense of fear or dread, as if something terrible is about to happen, despite there being no logical reason for it and interfere it with your ability to function normally in your daily life?",
            0: "I dont think so",
            1: "Yes, A couple of days",
            2: "The majority of the day, I am in this state.",
            3: "Yes, Every single day.",
        },
        {
            "question": "Its important to remember that we all make mistakes and experience setbacks, and that its possible to learn and grow from these experiences. Do you feel having difficulty focusing on certain things?",
            0: "Nope, not at all.",
            1: "Once in a while",
            2: "Yes, almost the whole day.",
            3: "Yes, Practically in every day.",
        },
        {
            "question": "Well, just remember that it's okay to make mistake and learn form it. Is there any chance that you have thoughts of harming yourself or feeling like you would be better off not living? If so better to address these thoughts as soon as possible and seek professional help, such as talking to a therapist or counselor, to get the support and guidance you need. Remember, you are not alone and there is help available to you.",
            0: "No, I never think that way.",
            1: "Somedays, it taking part of my mind",
            2: "It seems like I'm feeling this way all day long.",
            3: "I think im just anxious so I think this kind of stuff.",
        },
    ]

    answers = get_answers(phq9, input_list)
    print(answers)

    for answer in answers:
        print(answer)

    user = current_user
    user = User.query.filter_by(id=user.id).first()
    user.q1 = answers[0]
    user.q2 = answers[1]
    user.q3 = answers[2]
    user.q4 = answers[3]
    user.q5 = answers[4]
    user.q6 = answers[5]
    user.q7 = answers[6]
    user.q8 = answers[7]
    user.q9 = answers[8]
    user.result = int(input_list[9])
    user.epds_score = int(pred_result)
    user.date_submitted = datetime.now()
    # health = User_health(firstname=user.firstname, middlename=user.middlename,
    #                      lastname=user.lastname, address=user.address, contact=user.contact, condition=int(pred_result))
    # db.session.add(health)

    ppd = ""
    epds_class = ""
    if (pred_result) == 0:
        ppd = """I'm so glad that you're not feeling as down. It's always to have a positive outlook. And, Based on the conversation with me, You are likely to have a Minimal Risk or none mental health problem or possible developing a mental health disorder. If something doesn't seem quite right, it's important to start the conversation or consultation about getting help."""
        epds_class = "Minimal Risk"
    elif (pred_result) == 1:
        ppd = """I'm sorry to hear that but based on the conversation with me, You are possible to have a Low Risk level on mental health problem, learning about developing symptoms, or early warning signs, and taking action can help to ensure prompt treatment. Early intervention can help reduce the severity of this problem and interruptions in quality of life and functions. It may even be possible to delay or prevent a major mental illness altogether. Most mental problem don't improve on their own, and if untreated, a mental illness may get worse over time and cause serious problems."""
        epds_class = "Low Risk"
    elif (pred_result) == 2:
        ppd = """I understand that you're feeling a bit wary about what you're facing. Based on the conversation with me, You are possible in a Mild risk level of mental health problem that is something to look out for! When concerned that a you may be is a change in behavior or the presence of entirely new behaviors. It can make you miserable and can cause problems in your daily life, such as at school or work or in relationships. In most cases, symptoms can be managed with a combination of medications and talk therapy (psychotherapy)."""
        epds_class = "Mild Risk"
    elif (pred_result) == 3:
        ppd = """It's understandable to feel anxious in many situation. Based on the conversation with me, I can see that you are in Moderate risk level,  the sharpest concern like new or changed behavior that is related to a painful event, loss, or change probably people who take their lives exhibit one or more warning signs. It can cause severe emotional, behavioral and physical health problems, so it is important to seek attention from a primary care provider or mental health professional as soon as possible."""
        epds_class = "Moderate Risk"
    elif (pred_result) == 4:
        ppd = """I'm so sorry to hear you're thoughts and difficulties but Based on the conversation with me, You are probable in High risk Level that when the symptoms are intense and persistent, causing significant distress or impairment in daily functioning. Reach out to your doctor or mental health proffesional who can help you find the right path forward. You don't have to face this alone and I'm here for you."""
        epds_class = "High Risk"
    else:
        ppd = 'Nothing'

    data = {
        "id": str(user.id),
        "firstname": f"{user.firstname} {user.middlename} {user.lastname}",
        "email": user.email,
        "address": user.address,
        "date_submitted": str(user.date_submitted),
        "epds_score": epds_class,
        "questions": ["QUESTIONS", "Have you been capable of finding humor and laughing about situations?", "Have you anticipated things with pleasure and excitement?", "Have you needlessly held yourself responsible when things didn't go well?", "Have you experienced anxiety or concern without a valid cause?", "Have you experienced fear or panic without a clear or justifiable reason?", "Have things been overwhelming you?", "Have you been so unhappy that you have experienced trouble sleeping?", "Have you experienced feelings of sadness or misery?", "Have you been so unhappy that you have shed tears?", "EPDS Score"],
        "message": ppd.replace("Thankyou for answering my initial assessment. ", "")
    }

    table = ["Have you been capable of finding humor and laughing about situations?", "Have you anticipated things with pleasure and excitement?", "Have you needlessly held yourself responsible when things didn't go well?", "Have you experienced anxiety or concern without a valid cause?",
             "Have you experienced fear or panic without a clear or justifiable reason?", "Have things been overwhelming you?", "Have you been so unhappy that you have experienced trouble sleeping?", "Have you experienced feelings of sadness or misery?", "Have you been so unhappy that you have shed tears?"]

    epds = [
        {
            "ANSWER": "",
            0: "As much as I always could",
            1: "Not quite so much now",
            2: "Definitely not so much now",
            3: "Hardly at all",
        },
        {
            "question": "Have you anticipated things with pleasure and excitement?",
            0: "As much as I ever did",
            1: "Rather less than I used to",
            2: "Definitely less than I used to",
            3: "Hardly at all",
        },
        {
            "question": "Have you needlessly held yourself responsible when things didn't go well?",
            3: "Yes, most of the time",
            2: "Yes, some of the time",
            1: "Not very often",
            0: "No, Never",
        },
        {
            "question": "Have you experienced anxiety or concern without a valid cause?",
            0: "No, not at all",
            1: "Hardly ever",
            2: "Yes, sometimes",
            3: "Yes, very often",
        },
        {
            "question": "Have you experienced fear or panic without a clear or justifiable reason?",
            3: "Yes, quite a lot",
            2: "Yes, sometimes",
            1: "No, not much",
            0: "No, not at all",
        },
        {
            "question": "Have things been overwhelming you?",
            3: "Yes, most of the time I haven't been able to cope",
            2: "Yes, sometimes I haven't been coping as well as usual",
            1: "No, most of the time I have coped quite well",
            0: "No, I have been coping as well as ever",
        },
        {
            "question": "Have you been so unhappy that you have experienced trouble sleeping?",
            3: "Yes, most of the time",
            2: "Yes, sometimes",
            1: "Not very often",
            0: "No, not at all",
        },
        {
            "question": "Have you experienced feelings of sadness or misery?",
            3: "Yes, most of the time",
            2: "Yes, quite often",
            1: "Not very often",
            0: "No, not at all",

        },
        {
            "question": "Have you been so unhappy that you have shed tears?",
            3: "Yes, most of the time",
            2: "Yes, quite often",
            1: "Only occasionally",
            0: "No, never",
        },
        {
            "question": "Have you been so unhappy that you have shed tears?",
            3: "Yes, most of the time",
            2: "Yes, quite often",
            1: "Only occasionally",
            0: "No, never",
        },
        {
            "question": "Have you experienced feelings of sadness or misery?",
            3: "Yes, most of the time",
            2: "Yes, quite often",
            1: "Not very often",
            0: "",
        }
    ]

    document = Document()

    table = document.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    datas = ["id", "firstname", "email",
             "address", "date_submitted", "epds_score"]

    table.cell(0, 0).text = data['id']
    table.cell(0, 1).text = data['firstname']
    table.cell(0, 2).text = data["email"]
    table.cell(1, 0).text = data["address"]
    table.cell(1, 1).text = data["date_submitted"]
    table.cell(1, 2).text = data["epds_score"]

    document.add_paragraph()

    # Add a table to the document with 2 columns
    table = document.add_table(rows=len(data['questions']), cols=2)
    table.style = 'Table Grid'

    # Access the cells of the first row
    input_list.insert(0, "ANSWER")
    for index, value in enumerate(data['questions']):
        cell_1 = table.cell(index, 0)
        cell_2 = table.cell(index, 1)

        # Add text to the cells
        cell_1.text = value
        cell_2.text = f"({input_list[index]}) {epds[index][input_list[index] if input_list[index] in [0,1,2,3,4] else 0]}"
        if index == 0:
            cell_2.text = "ANSWER"
        elif index == 11:
            cell_2.text = f"{input_list[index]}"

    # table.paragraph_format.space_after = 1

    p = document.add_paragraph(data["message"])
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    table = document.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "EPDS Score"
    table.cell(0, 1).text = "Interpretation"
    table.cell(0, 2).text = "Action"

    table.cell(1, 0).text = "Less than 8"
    table.cell(1, 1).text = "Depression not likely"
    table.cell(1, 2).text = "Continue support"
    table.cell(2, 0).text = "9-11"
    table.cell(2, 1).text = "Depression possible"
    table.cell(
        2, 2).text = "Support, re-screen in 2–4 weeks. Consider referral to primary care provide(PCP)."
    table.cell(3, 0).text = "12-13"
    table.cell(3, 1).text = "Fairly high possibility of depression"
    table.cell(3, 2).text = "Monitor, support and offer education. Refer to PCP."
    table.cell(4, 0).text = "14 and higher (positive screen)"
    table.cell(4, 1).text = "Probable depression"
    table.cell(
        4, 2).text = "Diagnostic assessment and treatment by PCP and/or specialist."
    table.cell(
        5, 0).text = "Positive score (1, 2 or 3) on question 10 (suicidality risk)"
    table.cell(5, 1).text = ""
    table.cell(5, 2).text = "Immediate referral is necessary for assessment of suicidal ideation. This is to ensure proper intervention and to consider factors such as plan, history, symptoms, and potential harm."

    # temp_file = BytesIO()
    # document.save(temp_file)
    # temp_file.seek(0)

    new_filename = f"{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
    # user.file = Content(temp_file.read(), save=True)
    # user.file = temp_file
    # user.file.url

    document.save('./sanitybot/AssessmentResult.docx')
    db.session.commit()

    response = {}

    return ppd
